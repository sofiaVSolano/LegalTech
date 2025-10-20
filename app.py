# app.py
from flask import Flask, render_template, jsonify, request
from flask import send_file, make_response
import glob
import os
import uuid
import logging
from datetime import datetime
import json
from flask_socketio import SocketIO, emit
import openai
import base64
from io import BytesIO
import soundfile as sf
import numpy as np
from dotenv import load_dotenv 



# Inicializar la aplicación Flask
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key')

# Configurar Socket.IO para comunicación en tiempo real
socketio = SocketIO(app, cors_allowed_origins="*", logger=False, engineio_logger=False)

# Cargar variables de entorno PRIMERO, antes de cualquier otra cosa
load_dotenv()

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# Inicializar cliente OpenAI (API >=1.0.0)
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY no está configurada. Algunas funciones no estarán disponibles.")
else:
    logger.info(f"OPENAI_API_KEY cargada correctamente (longitud: {len(OPENAI_API_KEY)})")
client = openai.OpenAI(api_key=OPENAI_API_KEY)

# Estado de conversaciones en memoria
# En producción, esto debería almacenarse en una base de datos
conversations = {}

# Flujo de preguntas estructuradas para el asistente legal
QUESTIONS = [
    "Hola, ¿cuál es tu nombre completo?",
    "¿Cuál es tu número de cédula?",
    "¿Cuál es tu dirección y ciudad?",
    "Cuéntame tu historial legal.",
    # Eliminada la pregunta final para soportar modo de voz continuo (el usuario detiene cuando quiera)
]

@app.route('/')
def index():
    """Ruta principal que sirve la página de inicio del asistente legal."""
    return render_template('index.html')

@app.route('/api/start', methods=['POST'])
def start_conversation():
    """
    Inicia una nueva conversación con el asistente legal.
    Retorna el primer mensaje del asistente y un ID de conversación único.
    """
    try:
        # Generar un ID único para la conversación
        conv_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Guardar también un inicio en formato ISO para cálculos precisos de duración
        start_iso = datetime.now().isoformat()

        # Crear estado inicial de la conversación
        conversations[conv_id] = {
            'id': conv_id,
            'timestamp': timestamp,          # usado para nombres de archivo legibles
            'start_time': start_iso,         # ISO-8601, usado para cálculo de duración
            'current_question': 0,
            'responses': [],
            'complete': False,
            'connected': False,
            'duration': 0
        }
        
        # Registrar inicio de conversación
        logger.info(f"Started new conversation: {conv_id}")
        
        # Crear la primera respuesta del asistente
        first_message = QUESTIONS[0]
        
        # Agregar mensaje del asistente al historial
        conversations[conv_id]['responses'].append({
            'role': 'assistant',
            'content': first_message,
            'timestamp': datetime.now().isoformat()
        })
        
        # Retornar respuesta al cliente
        return jsonify({
            'conversation_id': conv_id,
            'message': first_message,
            'timestamp': timestamp
        })
        
    except Exception as e:
        logger.error(f"Error starting conversation: {str(e)}")
        return jsonify({'error': 'No se pudo iniciar la conversación'}), 500

@app.route('/api/message', methods=['POST'])
def handle_message():
    """
    Maneja un mensaje del usuario, procesa la respuesta y devuelve la siguiente pregunta.
    """
    try:
        # Obtener datos del request
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No se recibieron datos'}), 400
            
        conv_id = data.get('conversation_id')
        user_message = data.get('message', '').strip()
        
        # Validar ID de conversación
        if not conv_id:
            return jsonify({'error': 'Falta ID de conversación'}), 400

        # Log para depuración: mostrar lo que llegó y las conversaciones activas
        logger.info(f"/api/message received data: {data}")
        logger.debug(f"Conversations keys: {list(conversations.keys())}")

        # Obtener conversación
        conv = conversations.get(conv_id)
        if not conv:
            logger.warning(f"Conversation not found for id: {conv_id}")
            return jsonify({'error': 'Conversación no encontrada', 'conversation_id': conv_id, 'active_conversations': len(conversations)}), 404
        
        # Registrar mensaje del usuario
        if user_message:
            conv['responses'].append({
                'role': 'user',
                'content': user_message,
                'timestamp': datetime.now().isoformat()
            })
        
        # Determinar la siguiente acción basada en el estado actual
        current_q = conv['current_question']
        
        # Verificar si es la última pregunta y el usuario indica que no hay más información
        if current_q == len(QUESTIONS) - 1:
            # Normalizar la respuesta para comparación
            normalized_response = user_message.lower().strip()
            if any(neg in normalized_response for neg in ['no', 'nada', 'nada más', 'no hay', 'no quiero', 'no gracias']):
                conv['complete'] = True

                # Calcular duración usando start_time (ISO) cuando esté disponible.
                # Si no está disponible, intentar parsear el formato antiguo 'YYYYmmdd_HHMMSS'.
                start_dt = None
                start_iso = conv.get('start_time')
                if start_iso:
                    try:
                        start_dt = datetime.fromisoformat(start_iso)
                    except Exception:
                        start_dt = None

                if not start_dt:
                    # Fallback: intentar parsear 'timestamp' con el formato previo
                    try:
                        start_dt = datetime.strptime(conv.get('timestamp', ''), "%Y%m%d_%H%M%S")
                    except Exception:
                        start_dt = None

                if start_dt:
                    conv['duration'] = (datetime.now() - start_dt).total_seconds()
                else:
                    conv['duration'] = 0
                
                # Guardar la conversación
                save_conversation(conv_id)
                
                # Retornar mensaje de finalización
                return jsonify({
                    'message': 'Perfecto. He guardado la conversación.',
                    'conversation_complete': True,
                    'summary': {
                        'id': conv_id,
                        'timestamp': conv['timestamp'],
                        'duration': conv['duration'],
                        'questions_answered': len([r for r in conv['responses'] if r['role'] == 'user'])
                    }
                })
            else:
                # Si es la última pregunta pero la respuesta no es un 'no',
                # pedir explícitamente si desea agregar algo más.
                follow_up = '¿Hay algo más que quieras agregar?'
                conv['responses'].append({
                    'role': 'assistant',
                    'content': follow_up,
                    'timestamp': datetime.now().isoformat()
                })

                return jsonify({
                    'message': follow_up,
                    'conversation_complete': False,
                    'progress': {
                        'current': current_q + 1,
                        'total': len(QUESTIONS),
                        'percentage': int((current_q + 1) / len(QUESTIONS) * 100)
                    }
                })
        
        # Avanzar al siguiente estado si no estamos en la primera pregunta
        if current_q < len(QUESTIONS) - 1:
            conv['current_question'] += 1
        
        # Obtener la siguiente pregunta
        next_question_idx = conv['current_question']
        next_question = QUESTIONS[next_question_idx] if next_question_idx < len(QUESTIONS) else "Gracias por proporcionar la información."
        
        # Agregar la pregunta del asistente al historial
        conv['responses'].append({
            'role': 'assistant',
            'content': next_question,
            'timestamp': datetime.now().isoformat(),
            'question_number': next_question_idx + 1,
            'total_questions': len(QUESTIONS)
        })
        
        # Retornar la siguiente pregunta
        return jsonify({
            'message': next_question,
            'conversation_complete': False,
            'progress': {
                'current': next_question_idx + 1,
                'total': len(QUESTIONS),
                'percentage': int((next_question_idx + 1) / len(QUESTIONS) * 100)
            }
        })
        
    except Exception as e:
        logger.error(f"Error handling message: {str(e)}")
        return jsonify({'error': 'No se pudo procesar el mensaje'}), 500

@app.route('/api/save_manual', methods=['POST'])
def save_manual():
    """
    Permite guardar manualmente una conversación en curso.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No se recibieron datos'}), 400
            
        conv_id = data.get('conversation_id')
        
        # Validar ID de conversación
        if not conv_id:
            return jsonify({'error': 'Falta ID de conversación'}), 400
            
        if conv_id not in conversations:
            logger.warning(f"Conversation not found for id: {conv_id}")
            return jsonify({'error': 'Conversación no encontrada', 'conversation_id': conv_id, 'active_conversations': len(conversations)}), 404
        
        # Guardar la conversación
        save_conversation(conv_id)
        
        # Registrar acción
        logger.info(f"Manual save requested for conversation: {conv_id}")
        
        # Retornar confirmación
        return jsonify({
            'message': 'Conversación guardada manualmente.',
            'conversation_id': conv_id,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Error saving conversation manually: {str(e)}")
        return jsonify({'error': 'No se pudo guardar la conversación'}), 500

@app.route('/api/status', methods=['GET'])
def get_status():
    """
    Retorna el estado actual del servidor y las conversaciones activas.
    """
    try:
        active_conversations = len([cid for cid, conv in conversations.items() if not conv['complete']])
        
        status_data = {
            'server': 'active',
            'timestamp': datetime.now().isoformat(),
            'active_conversations': active_conversations,
            'total_conversations': len(conversations),
            'version': '1.0.0',
            'features': {
                'voice_recognition': True,
                'text_to_speech': True,
                'openai_realtime': bool(OPENAI_API_KEY),
                'manual_save': True
            }
        }
        
        return jsonify(status_data)
        
    except Exception as e:
        logger.error(f"Error getting status: {str(e)}")
        return jsonify({'error': 'No se pudo obtener el estado'}), 500


def _extract_name_from_conv(conv):
    """Intentar extraer el nombre de la persona desde las respuestas guardadas."""
    try:
        # Buscar la primera respuesta del usuario
        for r in conv.get('responses', []):
            if r.get('role') == 'user' and r.get('content'):
                # Tomar la primera respuesta de usuario como nombre (mejor heurística que nada)
                candidate = r.get('content').strip()
                # limitar longitud
                return candidate if len(candidate) < 120 else candidate[:120]
    except Exception:
        pass
    return 'Sin nombre'


@app.route('/conversaciones')
def conversaciones_page():
    """Página que lista las conversaciones guardadas (desde memoria)."""
    return render_template('conversations.html')


@app.route('/api/conversations/list', methods=['GET'])
def api_conversations_list():
    try:
        items = []
        for cid, conv in conversations.items():
            items.append({
                'id': cid,
                'name': _extract_name_from_conv(conv),
                'timestamp': conv.get('timestamp'),
                'complete': conv.get('complete', False)
            })

        # También incluir conversaciones ya guardadas en disco que no estén en memoria
        files = glob.glob('conversaciones/conversacion_*.json')
        for f in files:
            try:
                with open(f, 'r', encoding='utf-8') as fh:
                    conv = json.load(fh)
                    cid = conv.get('id')
                    if cid not in conversations:
                        items.append({
                            'id': cid,
                            'name': _extract_name_from_conv(conv),
                            'timestamp': conv.get('timestamp'),
                            'complete': conv.get('complete', False)
                        })
            except Exception:
                continue

        # Ordenar por timestamp descendente
        items.sort(key=lambda x: x.get('timestamp') or '', reverse=True)
        return jsonify({'conversations': items})
    except Exception as e:
        logger.error(f"Error listing conversations: {str(e)}")
        return jsonify({'error': 'No se pudo listar las conversaciones'}), 500


@app.route('/conversacion/<conv_id>')
def conversation_view_page(conv_id):
    return render_template('conversation_view.html', conversation_id=conv_id)


@app.route('/api/conversation/<conv_id>', methods=['GET'])
def api_get_conversation(conv_id):
    try:
        conv = conversations.get(conv_id)
        if not conv:
            # intentar cargar desde disco
            matches = glob.glob(f'conversaciones/conversacion_{conv_id}_*.json')
            if matches:
                with open(matches[0], 'r', encoding='utf-8') as fh:
                    conv = json.load(fh)
            else:
                return jsonify({'error': 'Conversación no encontrada'}), 404
        return jsonify({'conversation': conv})
    except Exception as e:
        logger.error(f"Error getting conversation {conv_id}: {str(e)}")
        return jsonify({'error': 'No se pudo obtener la conversación'}), 500


@app.route('/api/conversation/<conv_id>/delete', methods=['POST'])
def api_delete_conversation(conv_id):
    try:
        # eliminar de memoria
        if conv_id in conversations:
            del conversations[conv_id]

        # eliminar archivos en disco que coincidan
        patterns = [f'conversaciones/conversacion_{conv_id}_*.json', f'conversaciones/conversacion_{conv_id}_*.txt']
        removed = []
        for p in patterns:
            for f in glob.glob(p):
                try:
                    os.remove(f)
                    removed.append(f)
                except Exception:
                    logger.warning(f"No se pudo eliminar archivo: {f}")

        return jsonify({'deleted_files': removed, 'conversation_id': conv_id})
    except Exception as e:
        logger.error(f"Error deleting conversation {conv_id}: {str(e)}")
        return jsonify({'error': 'No se pudo eliminar la conversación'}), 500


@app.route('/api/conversation/<conv_id>/entender', methods=['POST'])
def api_conversation_entender(conv_id):
    """Procesa la conversación con OpenAI para redactar la carta formal y clasificar el caso."""
    try:
        conv = conversations.get(conv_id)
        if not conv:
            # intentar cargar desde disco
            matches = glob.glob(f'conversaciones/conversacion_{conv_id}_*.json')
            if matches:
                with open(matches[0], 'r', encoding='utf-8') as fh:
                    conv = json.load(fh)
            else:
                return jsonify({'error': 'Conversación no encontrada'}), 404

        # Preparar prompt con todo el histórico de la conversación
        convo_text = ''
        for r in conv.get('responses', []):
            role = r.get('role', '')
            content = r.get('content', '')
            convo_text += f"{role.upper()}: {content}\n"

        system_msg = (
            "Eres un abogado experto y especializado. Lee la conversación proporcionada y: "
            "1) Redacta una carta formal exponiendo el caso del cliente explicándolo con el tecnisismo adecuado"
            "2) Expón claramente los pasos recomendados para resolver el problema. "
            "3) Da una estimación aproximada del costo en pesos colombianos (rango). "
            "4) Clasifica el asunto en UNA de las siguientes categorías: Derecho Público (con subdivisiones), Derecho Privado (con subdivisiones), o Derecho Social (con subdivisiones). "
            "Responde únicamente en JSON con las claves: category, subdivision, letter, recommendations, estimated_cost. "
            "Usa español formal."
        )

        user_msg = f"Conversación:\n{convo_text}\nPor favor devuelve el JSON requerido." 

        # Verificar que tenemos la API key configurada
        if not OPENAI_API_KEY:
            logger.error('OPENAI_API_KEY no configurada; no se puede procesar con OpenAI')
            return jsonify({'error': 'OPENAI_API_KEY no está configurada en el servidor. Configure la variable de entorno.'}), 503

        # Intentar llamar al modelo (usar GPT_MODEL si está definido)
        model = os.environ.get('GPT_MODEL', 'gpt-4-1106-preview')
        resp = None
        assistant_text = None
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {'role': 'system', 'content': system_msg},
                    {'role': 'user', 'content': user_msg}
                ],
                max_tokens=1200,
                temperature=0.6
            )
        except Exception as e_first:
            logger.warning(f"Fallo llamando a {model}: {e_first}. Intentando con gpt-4")
            try:
                resp = client.chat.completions.create(
                    model='gpt-4',
                    messages=[
                        {'role': 'system', 'content': system_msg},
                        {'role': 'user', 'content': user_msg}
                    ],
                    max_tokens=1200,
                    temperature=0.3
                )
            except Exception as e_second:
                # Log completo y retornar error con detalles para debugging
                logger.exception(f"Error llamando a OpenAI (intentos a {model} y gpt-4 fallaron)")
                return jsonify({'error': 'Error llamando a OpenAI', 'details': str(e_second)}), 502

        # Extraer texto de la respuesta de forma robusta
        try:
            # Intentar acceder a la estructura esperada
            if hasattr(resp, 'choices') and len(resp.choices) > 0:
                # Manejar objetos tipo dot-access
                first = resp.choices[0]
                if hasattr(first, 'message') and hasattr(first.message, 'content'):
                    assistant_text = first.message.content
                elif isinstance(first, dict) and 'message' in first and 'content' in first['message']:
                    assistant_text = first['message']['content']
            # Fallback a str(resp)
            if assistant_text is None:
                assistant_text = str(resp)
        except Exception:
            logger.exception('No se pudo extraer texto de la respuesta de OpenAI')
            assistant_text = str(resp)
        # Intentar parsear JSON resultante
        parsed = None
        try:
            # Buscar primer bloque JSON en el texto
            import re
            m = re.search(r"(\{[\s\S]*\})", assistant_text)
            if m:
                parsed = json.loads(m.group(1))
        except Exception:
            logger.warning('No se pudo parsear JSON desde la respuesta de OpenAI')

        # Asegurar que la conversación está en memoria para poder persistir el análisis
        if conv_id not in conversations:
            conversations[conv_id] = conv

        # Guardar el análisis en la conversación y persistir en disco
        if parsed:
            conversations[conv_id].setdefault('analysis', {})
            conversations[conv_id]['analysis'].update(parsed)
            conversations[conv_id]['analysis_raw'] = assistant_text
        else:
            conversations[conv_id].setdefault('analysis', {})
            conversations[conv_id]['analysis_raw'] = assistant_text

        # Guardar en disco
        save_conversation(conv_id)

        # Retornar el resultado (JSON parsed si está disponible, sino texto)
        if parsed:
            return jsonify({'result': parsed})
        return jsonify({'result_text': assistant_text})

    except Exception as e:
        logger.error(f"Error processing 'entender' for {conv_id}: {str(e)}")
        return jsonify({'error': 'No se pudo procesar la conversación con OpenAI'}), 500


@app.route('/api/conversation/<conv_id>/download.txt', methods=['GET'])
def api_download_letter_txt(conv_id):
    try:
        conv = conversations.get(conv_id)
        if not conv:
            # intentar cargar desde disco
            matches = glob.glob(f'conversaciones/conversacion_{conv_id}_*.json')
            if matches:
                with open(matches[0], 'r', encoding='utf-8') as fh:
                    conv = json.load(fh)
            else:
                return jsonify({'error': 'Conversación no encontrada'}), 404

        analysis = conv.get('analysis') or {}
        letter = analysis.get('letter') or conv.get('analysis_raw') or 'No hay carta disponible.'

        # Preparar respuesta como archivo de texto
        resp = make_response(letter)
        resp.headers['Content-Type'] = 'text/plain; charset=utf-8'
        resp.headers['Content-Disposition'] = f'attachment; filename=letter_{conv_id}.txt'
        return resp
    except Exception as e:
        logger.error(f"Error downloading txt for {conv_id}: {e}")
        return jsonify({'error': 'No se pudo generar el archivo txt'}), 500


@app.route('/api/conversation/<conv_id>/download.docx', methods=['GET'])
def api_download_letter_docx(conv_id):
    try:
        conv = conversations.get(conv_id)
        if not conv:
            matches = glob.glob(f'conversaciones/conversacion_{conv_id}_*.json')
            if matches:
                with open(matches[0], 'r', encoding='utf-8') as fh:
                    conv = json.load(fh)
            else:
                return jsonify({'error': 'Conversación no encontrada'}), 404

        analysis = conv.get('analysis') or {}
        letter = analysis.get('letter') or conv.get('analysis_raw') or 'No hay carta disponible.'

        # Generar .docx en memoria
        try:
            from docx import Document
        except Exception:
            return jsonify({'error': 'python-docx no está disponible en el entorno'}), 500

        doc = Document()
        # Añadir título
        doc.add_heading('Carta formal - Asistente Legal', level=1)
        # Añadir metadata básica
        doc.add_paragraph(f'ID conversación: {conv_id}')
        doc.add_paragraph('')

        # Añadir el contenido de la carta, preservando saltos de línea
        for line in letter.split('\n'):
            doc.add_paragraph(line)

        # Guardar en un BytesIO
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        return send_file(bio, as_attachment=True, download_name=f'letter_{conv_id}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        logger.error(f"Error generating docx for {conv_id}: {e}")
        return jsonify({'error': 'No se pudo generar el archivo docx'}), 500

def save_conversation(conv_id):
    """
    Guarda una conversación en formato JSON y texto plano.
    """
    try:
        if conv_id not in conversations:
            logger.warning(f"Attempt to save non-existent conversation: {conv_id}")
            return False
        
        conv = conversations[conv_id]
        filename = f"conversaciones/conversacion_{conv_id}_{conv['timestamp']}"
        
        # Crear directorio si no existe
        os.makedirs('conversaciones', exist_ok=True)
        
        # Guardar en formato JSON
        json_path = f"{filename}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(conv, f, ensure_ascii=False, indent=2)
        
        # Guardar en formato texto plano
        txt_path = f"{filename}.txt"
        with open(txt_path, 'w', encoding='utf-8') as f:
            # Preferir mostrar start_time (ISO) si está disponible para mayor claridad
            header_time = conv.get('start_time') or conv.get('timestamp')
            f.write(f"CONVERSACIÓN LEGAL - {header_time}\n")
            f.write(f"ID: {conv['id']}\n")
            f.write(f"Estado: {'Completada' if conv['complete'] else 'Guardada manualmente'}\n")
            if 'duration' in conv and isinstance(conv.get('duration'), (int, float)):
                f.write(f"Duración: {conv['duration']:.0f} segundos\n")
            f.write("-" * 50 + "\n\n")
            
            for response in conv['responses']:
                f.write(f"{response['role'].capitalize()}: {response['content']}\n")
        
        logger.info(f"Conversation saved: {filename}")
        return True
        
    except Exception as e:
        logger.error(f"Error saving conversation {conv_id}: {str(e)}")
        return False

# Manejadores de eventos Socket.IO para comunicación en tiempo real
@socketio.on('connect')
def handle_connect():
    """
    Maneja la conexión de un cliente a través de Socket.IO.
    """
    logger.info(f"Client connected: {request.sid}")
    emit('connection_status', {'status': 'connected', 'sid': request.sid})

@socketio.on('disconnect')
def handle_disconnect():
    """
    Maneja la desconexión de un cliente.
    """
    logger.info(f"Client disconnected: {request.sid}")

@socketio.on('openai_connect')
def handle_openai_connect(data):
    """
    Maneja la solicitud de conexión con la API en tiempo real de OpenAI.
    """
    try:
        conv_id = data.get('conversation_id')
        
        if not conv_id or conv_id not in conversations:
            emit('error', {'message': 'Conversación no válida'})
            return
        
        # Actualizar estado de la conversación
        conversations[conv_id]['connected'] = True
        conversations[conv_id]['realtime_start'] = datetime.now().isoformat()
        
        logger.info(f"OpenAI Realtime connection requested for conversation: {conv_id}")
        
        # Emitir confirmación de conexión
        emit('openai_connected', {
            'status': 'connected', 
            'sid': request.sid,
            'conversation_id': conv_id,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Error in openai_connect: {str(e)}")
        emit('error', {'message': 'No se pudo conectar con OpenAI Realtime'})

@socketio.on('audio_data')
def handle_audio_data(data):
    """
    Maneja datos de audio enviados por el cliente.
    """
    try:
        conv_id = data.get('conversation_id')
        audio_data = data.get('audio')
        
        if not conv_id or conv_id not in conversations:
            emit('error', {'message': 'Conversación no válida'})
            return
            
        if not audio_data:
            emit('error', {'message': 'No se recibieron datos de audio'})
            return
        
        # Registrar recepción de audio
        logger.info(f"Audio data received for conversation {conv_id}: {len(audio_data)} characters")
        
        # En una implementación completa, aquí se procesaría el audio
        # y se enviaría a OpenAI Realtime API
        
        # Para esta demo, simulamos una respuesta
        emit('openai_response', {
            'text': 'Estoy procesando tu solicitud. ¿Puedes proporcionar más detalles sobre tu caso legal?',
            'conversation_id': conv_id,
            'timestamp': datetime.now().isoformat(),
            'processing_time': 0.5
        })
        
    except Exception as e:
        logger.error(f"Error processing audio data: {str(e)}")
        emit('error', {'message': 'Error procesando audio'})


@app.route('/api/upload_voice_response', methods=['POST'])
def upload_voice_response():
    """
    Endpoint para recibir archivos de audio (grabaciones manuales) desde el cliente.
    Guarda el archivo en el directorio `conversaciones` y lo asocia a la conversación en memoria si existe.
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No se encontró el archivo'}), 400

        file = request.files['file']
        conv_id = request.form.get('conversation_id') or request.args.get('conversation_id')

        if not conv_id:
            return jsonify({'error': 'Falta conversation_id'}), 400

        # Crear directorio si no existe
        os.makedirs('conversaciones', exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'conversaciones/audio_{conv_id}_{timestamp}_{uuid.uuid4().hex}.webm'

        # Guardar archivo
        file.save(filename)
        logger.info(f'Audio uploaded and saved: {filename}')

        # Asociar referencia en la conversación en memoria (si existe)
        if conv_id in conversations:
            conversations[conv_id].setdefault('responses', [])
            conversations[conv_id]['responses'].append({
                'role': 'user',
                'content': f'[Adjunto audio] {filename}',
                'timestamp': datetime.now().isoformat(),
                'type': 'audio',
                'file': filename
            })

        return jsonify({'message': 'Audio recibido y guardado', 'path': filename})

    except Exception as e:
        logger.error(f'Error uploading voice response: {str(e)}')
        return jsonify({'error': 'Error al subir el audio'}), 500

@socketio.on('text_input')
def handle_text_input(data):
    """
    Maneja entrada de texto para procesamiento con OpenAI.
    """
    try:
        conv_id = data.get('conversation_id')
        text = data.get('text')
        
        if not conv_id or conv_id not in conversations:
            emit('error', {'message': 'Conversación no válida'})
            return
            
        if not text:
            emit('error', {'message': 'No se recibió texto'})
            return
        
        # Registrar entrada de texto
        logger.info(f"Text input received for conversation {conv_id}: {text[:50]}...")
        
        # Simular procesamiento con OpenAI
        # En una implementación completa, aquí se llamaría a la API de OpenAI
        response_text = f"He entendido que tu caso legal involucra: {text}. ¿Puedes proporcionar más detalles específicos sobre las partes involucradas?"
        
        # Emitir respuesta
        emit('openai_response', {
            'text': response_text,
            'conversation_id': conv_id,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Error processing text input: {str(e)}")
        emit('error', {'message': 'Error procesando texto'})

# Manejador de errores global
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Ruta no encontrada'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Error interno del servidor'}), 500

# Punto de entrada principal
if __name__ == '__main__':
    """
    Punto de entrada para ejecutar la aplicación.
    Crea los directorios necesarios y inicia el servidor.
    """
    # Crear directorio de conversaciones si no existe
    if not os.path.exists('conversaciones'):
        os.makedirs('conversaciones')
        logger.info("Created 'conversaciones' directory")

    # Iniciar el servidor Socket.IO con eventlet para soporte WebSocket
    logger.info("Starting Flask server with Socket.IO (eventlet)...")
    logger.info("Application is ready. Access at http://localhost:5000")
    import eventlet
    import eventlet.wsgi
    socketio.run(
        app,
        host='0.0.0.0',
        port=int(os.environ.get('PORT', 5000)),
        debug=os.environ.get('FLASK_ENV') == 'development',
        use_reloader=os.environ.get('FLASK_ENV') == 'development'
    )
